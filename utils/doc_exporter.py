from docx import Document
from docx.shared import RGBColor
from io import BytesIO
import re


def add_bold_runs(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)

        # Force black color
        run.font.color.rgb = RGBColor(0, 0, 0)


def markdown_to_docx(markdown_text: str):
    document = Document()
    lines = markdown_text.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            heading = document.add_heading(level=1)
            add_bold_runs(heading, line.replace("# ", ""))

        elif line.startswith("## "):
            heading = document.add_heading(level=2)
            add_bold_runs(heading, line.replace("## ", ""))

        elif line.startswith("### "):
            heading = document.add_heading(level=3)
            add_bold_runs(heading, line.replace("### ", ""))

        elif line.startswith("- ") or line.startswith("• "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_bold_runs(paragraph, line[2:])

        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            paragraph = document.add_paragraph(style="List Number")
            add_bold_runs(paragraph, text)

        else:
            paragraph = document.add_paragraph()
            add_bold_runs(paragraph, line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer
